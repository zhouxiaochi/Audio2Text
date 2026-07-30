import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from backend.models import JobRecord, TranscriptSegment


def timestamp(seconds: float) -> str:
    total = max(0, round(seconds))
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def generate_markdown(job: JobRecord, segments: list[TranscriptSegment]) -> str:
    lines = [
        f"# {job.original_filename}",
        "",
        f"- 任务 ID：`{job.id}`",
        f"- 生成时间：{job.updated_at}",
        "",
        "## 转录与翻译",
        "",
    ]
    for segment in segments:
        speaker = segment.speaker or "Speaker"
        lines.extend(
            [
                f"### [{timestamp(segment.start)}–{timestamp(segment.end)}] {speaker}",
                "",
                segment.text,
                "",
                segment.translation_zh or "",
                "",
            ]
        )
    return "\n".join(lines).rstrip() + "\n"


def write_artifacts(
    job_dir: Path, job: JobRecord, segments: list[TranscriptSegment]
) -> dict[str, Path]:
    job_dir.mkdir(parents=True, exist_ok=True)
    json_path = job_dir / "transcript.json"
    markdown_path = job_dir / "transcript.md"
    json_path.write_text(
        json.dumps([item.model_dump() for item in segments], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    markdown_path.write_text(generate_markdown(job, segments), encoding="utf-8")
    docx_path = job_dir / "transcript.docx"
    render_docx(markdown_path.read_text(encoding="utf-8"), docx_path)
    return {"json": json_path, "md": markdown_path, "docx": docx_path}


def render_docx(markdown: str, destination: Path) -> None:
    """Render the generated Markdown subset into a standalone DOCX document."""

    document = Document()
    styles = document.styles
    # Noto Sans CJK is available in the Linux deployment image and remains
    # readable when Word substitutes the font on another operating system.
    styles["Normal"].font.name = "Noto Sans CJK SC"
    styles["Normal"].font.size = Pt(10.5)
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            document.add_heading(heading.group(2), level=len(heading.group(1)))
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
